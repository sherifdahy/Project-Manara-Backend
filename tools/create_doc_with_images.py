import os
import sys
from pathlib import Path

try:
    import cairosvg
    from docx import Document
    from docx.shared import Inches
except Exception as e:
    print("Missing dependencies:", e)
    print("Run: python -m pip install cairosvg python-docx pillow")
    sys.exit(2)

ROOT = Path(__file__).resolve().parents[1]
GUIDE_TXT = ROOT / 'DESIGN_PATTERNS_COMPREHENSIVE_GUIDE.txt'
DIAGRAMS_DIR = ROOT / 'diagrams'
OUTPUT_DOCX = ROOT / 'DESIGN_PATTERNS_COMPREHENSIVE_GUIDE.docx'

if not GUIDE_TXT.exists():
    print(f"Guide text not found: {GUIDE_TXT}")
    sys.exit(1)

svg_files = sorted([p for p in DIAGRAMS_DIR.glob('*.svg')])
if not svg_files:
    print(f"No SVG files found in {DIAGRAMS_DIR}")

# Ensure diagrams dir exists
if not DIAGRAMS_DIR.exists():
    print(f"Diagrams directory missing: {DIAGRAMS_DIR}")
    sys.exit(1)

# Create PNGs list
png_files = []
for svg in svg_files:
    png_path = svg.with_suffix('.png')
    try:
        cairosvg.svg2png(url=str(svg), write_to=str(png_path))
        png_files.append(png_path)
        print(f"Converted: {svg.name} -> {png_path.name}")
    except Exception as ex:
        print(f"Failed to convert {svg}: {ex}")

# Build DOCX
doc = Document()
doc.add_heading('Design Patterns Comprehensive Guide', level=1)

with GUIDE_TXT.open('r', encoding='utf-8') as f:
    content = f.read()

# Split into paragraphs by double newlines to keep formatting
for para in content.split('\n\n'):
    doc.add_paragraph(para.strip())

# Add images with captions
if png_files:
    doc.add_page_break()
    doc.add_heading('Diagrams', level=2)
    for png in png_files:
        try:
            # Insert image, scale to width 6 inches max
            doc.add_paragraph(png.stem)
            doc.add_picture(str(png), width=Inches(6))
        except Exception as ex:
            print(f"Failed to insert image {png}: {ex}")

# Save
try:
    doc.save(str(OUTPUT_DOCX))
    print(f"Created DOCX: {OUTPUT_DOCX}")
except Exception as ex:
    print(f"Failed to save DOCX: {ex}")
    sys.exit(1)
