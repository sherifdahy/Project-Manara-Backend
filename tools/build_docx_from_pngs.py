import re
from pathlib import Path
from docx import Document
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[1]
GUIDE = ROOT / 'DESIGN_PATTERNS_COMPREHENSIVE_GUIDE.txt'
DIAGRAMS = ROOT / 'diagrams'
OUT = ROOT / 'DESIGN_PATTERNS_COMPREHENSIVE_GUIDE.docx'

if not GUIDE.exists():
    print('Guide file not found:', GUIDE)
    raise SystemExit(1)

with GUIDE.open('r', encoding='utf-8') as f:
    lines = f.readlines()

# Remove ASCII-art blocks starting/ending with lines that contain +--- sequences
clean_lines = []
skip = False
for i, line in enumerate(lines):
    if not skip and re.search(r"\+[-=]{3,}", line):
        skip = True
        continue
    if skip and re.search(r"\+[-=]{3,}", line):
        skip = False
        continue
    if not skip:
        clean_lines.append(line)

content = ''.join(clean_lines).strip()

# Build DOCX
doc = Document()
doc.add_heading('Design Patterns Comprehensive Guide', level=1)

for para in content.split('\n\n'):
    text = para.strip()
    if not text:
        continue
    doc.add_paragraph(text)

# Insert diagrams (PNG files) in a sensible order if present
pngs = sorted([p for p in DIAGRAMS.glob('*.png')])
if pngs:
    doc.add_page_break()
    doc.add_heading('Diagrams', level=2)
    for png in pngs:
        doc.add_paragraph(png.stem)
        try:
            doc.add_picture(str(png), width=Inches(6))
        except Exception as e:
            print('Failed to add picture', png, e)

try:
    doc.save(str(OUT))
    print('Saved DOCX:', OUT)
except Exception as e:
    print('Failed to save DOCX:', e)
    raise
